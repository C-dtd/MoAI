
from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.chat_models import ChatOllama
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.prompts.prompt import PromptTemplate
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from docx import Document
from flask import *
from flask_cors import CORS
import json
import os
import sys
import datetime

app = Flask(__name__)
# CORS(app, resources={r'*': {'origins': 'http://localhost:8000'}})
CORS(app)
host = 'localhost'
port = 5100

# Configuration
ngrok = 'https://6c82-35-229-167-67.ngrok-free.app'  # Replace with your ngrok URL
device = 'cpu'

# Load language model
llm_model = ChatOllama(
    model='meta-llama-3.1',
    base_url=ngrok
)

# Load embeddings model
embedding_model = HuggingFaceEmbeddings(
    model_name="BAAI/bge-m3",
    model_kwargs={'device': device},
    encode_kwargs={'normalize_embeddings': True}
)

# Define prompt
prompt_template = '''Use the following pieces of context to answer the question at the end.
If you don't find the answer in context, don't try to make up an answer.
If you find the answer in context, answer me only use korean.

context: {context}

Question: {question}
Helpful Answer:'''
rag_prompt = PromptTemplate.from_template(prompt_template)

# Function to create the DOCX file
def create_docx(response, vectorstore):
    doc = Document()
    title = response.get('title', '제목 없음')
    doc.add_heading(title, level=0)

    doc.add_heading('목차', level=1)
    doc.add_paragraph('1. 개요')
    doc.add_paragraph('2. 본문')
    for n, cont in enumerate(response.get('content', []), 1):
        doc.add_paragraph(f'\t2-{n}. {cont}')

    doc.add_heading('1. 개요', level=1)
    doc.add_paragraph(response.get('summary', ''))

    doc.add_heading('2. 본문', level=1)
    for n, cont in enumerate(response.get('content', []), 1):
        question = f'{title}라는 보고서의 {cont} 부분 상세 내용을 글로 풀어서 적어줘'    # {title}라는 보고서의 {cont} 부분 상세 내용 markdown 형식으로
        memory = ConversationBufferMemory(
            memory_key='chat_history',
            return_messages=True
        )
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=llm_model,
            retriever=vectorstore.as_retriever(),
            condense_question_prompt=rag_prompt,
            memory=memory
        )
        res = conversation_chain({'question': question})
        doc.add_heading(f'\t2-{n}. {cont}', level=2)
        doc.add_paragraph(res['chat_history'][1].content)
    
    # Format date to avoid issues in filenames
    doc.save(f'./processed/{datetime.datetime.today().strftime("%g%m%d%H%M%S")}.docx')
    return f'/processed/{datetime.datetime.today().strftime("%g%m%d%H%M%S")}.docx'

def process_file(file_path):
    # if not os.path.isfile(file_path):
    #     print(f'파일이 존재하지 않습니다: {file_path}')
    #     return
    
    text_sum = ''
    # files = [file_path]
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    for file in file_path:
        reader = PdfReader(file)
        for page in reader.pages:  # 페이지 별로 텍스트 추출
            text = page.extract_text()
            corrected_text = text.encode('utf-8', errors='ignore').decode('utf-8')  # 인코딩 오류 무시 및 텍스트 누적
            text_sum += corrected_text + '\n'
    splits = text_splitter.split_text(text_sum)

    # Create FAISS index
    vectorstore = FAISS.from_texts(splits, embedding_model)
    
    # Assuming text processing and JSON generation happens here
    # For demonstration, we're creating a sample response
    response = {
        'title': '보고서 제목',
        'content': ['목차 1', '목차 2'],
        'summary': '보고서 개요'
    }

    return create_docx(response, vectorstore)

@app.route('/summary', methods=['POST'])
def summary():
    
    files = [request.files[i] for i in request.files]
    processedFilePath = process_file(files)
    return jsonify({
        'result': 'ok',
        'processedFilePath': processedFilePath
    })

@app.route('/embedding', methods=['POST'])
def embedding():
    params = request.json()

if __name__ == '__main__':
    app.run(host= host, port=port)
    
    # if len(sys.argv) != 2:
    #     print("Usage: python ollama.py <file_path>")
    # else:
    #     process_file(sys.argv[1])