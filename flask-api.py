from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.chat_models import ChatOllama
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.prompts.prompt import PromptTemplate
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from docx import Document
from flask import *
from flask_cors import CORS
import json
import os
import sys
import datetime
from langchain_postgres import PGVector
from langchain_postgres.vectorstores import PGVector
import psycopg2
from psycopg2 import pool

dbcp = psycopg2.pool.SimpleConnectionPool(1, 20,
    user='postgres.vpcdvbdktvvzrvjfyyzm',
    password='Odvv8E1iChKjwai4',
    host='aws-0-ap-southeast-1.pooler.supabase.com',
    port=6543,
    dbname='postgres'
)

app = Flask(__name__)
# CORS(app, resources={r'*': {'origins': 'http://localhost:8000'}})
CORS(app)
host = 'localhost'
port = 5100

connection='postgresql+psycopg2://postgres.vpcdvbdktvvzrvjfyyzm:Odvv8E1iChKjwai4@aws-0-ap-southeast-1.pooler.supabase.com:6543/postgres'

# Configuration
ngrok = ''
ngrok = 'https://b9b3-34-142-171-135.ngrok-free.app'
device = 'cpu'

if ngrok == '':
    llm_model = ChatOllama(
        model='meta-llama-3.1',
    )
    llm_model_json = ChatOllama(
        model='meta-llama-3.1',
        format='json', 
    )
else:
    # Load language model
    llm_model = ChatOllama(
        model='meta-llama-3.1',
        base_url=ngrok      # 주석 해제 시 코랩 자원으로 돌아감, # 주석 설정 시 로컬 자원으로 돌아감 
    )
    llm_model_json = ChatOllama(
        model='meta-llama-3.1',
        format='json', 
        base_url=ngrok       # 주석 해제 시 코랩 자원으로 돌아감, # 주석 설정 시 로컬 자원으로 돌아감 
    )

# Load embeddings model
embedding_model = HuggingFaceEmbeddings(
    model_name="BAAI/bge-m3",
    model_kwargs={'device': device},
    encode_kwargs={'normalize_embeddings': True}
)

text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

# Define prompt
prompt_template = '''Use the following pieces of context to answer the question at the end.
If you don't find the answer in context, don't try to make up an answer.
If you find the answer in context, answer me only use korean.

context: {chat_history}

Question: {question}
Helpful Answer:'''
# rag_prompt = PromptTemplate.from_template(prompt_template)
rag_prompt = PromptTemplate(template=prompt_template, input_variables=['chat_history', 'question'])

prompt_template = '''Use the following pieces of context to answer the question at the end.
If you don't find the answer in context, don't try to make up an answer.
If you find the answer in context, answer me only use korean.

context: {chat_history}

Question: {question}
Helpful Answer:'''
# rag_prompt = PromptTemplate.from_template(prompt_template)
chat_prompt = PromptTemplate(template=prompt_template, input_variables=['chat_history', 'question'])



#db에서 접근 권한이 있는 정보를 문장으로 만들어 분할하는 함수
def db_crawler(user_id):
    db = dbcp.getconn()
    cur = db.cursor()
    cur.execute("""
                select user_name, job_id, dep_name 
                from users u join departments d on u.dep_id = d.dep_id
                where user_id=%s
                """,
                (user_id,))
    user_info = cur.fetchone()
    cur.execute("""
                select cl.chat_at, u.user_name, cl.cl_chat
                from chat_logs cl join users u on cl.user_id = u.user_id
                where cl_type = 'text' and room_id in (select room_id from room_users where user_id = %s)
                order by chat_at
                """,
                (user_id,))
    chat_rows = cur.fetchall()
    cur.execute("""
                select cal_title, cal_start_date, cal_end_date, cal_location
                from calendars
                where user_id = %s or cal_id in (select cal_sub_id from calendar_shared)
                order by cal_start_date
                """,
                (user_id,))
    cal_rows = cur.fetchall()
    cur.execute("""
                select user_name, job_id, dep_name, user_phone
                from users u join departments d on u.dep_id = d.dep_id
                """)
    member_rows = cur.fetchall()
    cur.execute("""
                select c_title, c_category, c_date_range, c_assignee
                from cards
                where dep_id = (
                    select dep_id
                    from users
                    where user_id = %s
                )
                """,
                (user_id,))
    cur.execute("""
                select p_title, user_name, app_at
                from payment p join users u on p.p_app = u.user_id
                where p_uploader = %s
                """,
                (user_id,))
    payreq_rows = cur.fetchall()
    cur.execute("""
                select p_title, user_name, app_at
                from payment p join users u on p.p_uploader = u.user_id
                where p_app =  %s 
                """,
                (user_id,))
    payres_rows = cur.fetchall()
    task_rows = cur.fetchall()
    
    dbcp.putconn(db)

    db_crawls = f'사용자 이름: {user_info[0]}, 사용자 직책: {user_info[1]}, 사용자 부서: {user_info[2]}'
    db_crawls += datetime.datetime.now().strftime('오늘은 %G-%m-%d, %A입니다.')

    for row in chat_rows:
        _ = '\n'
        db_crawls += f"{row[0].strftime('%G-%m-%d %T')}에 {row[1]}이(가) '{row[2].replace(_, '')}'라 말함. "
    for row in cal_rows:
        db_crawls += f"{row[1].replace('T', ' ')}부터 {row[2].replace('T', ' ')}까지 일정: '{row[0]}'이 {row[3] +'에서 ' if row[3] else ''}있습니다. "
    for row in member_rows:--
        db_crawls += f"{row[0]}은(는) {row[2]} 부서의 {row[1]} 직책을 담당하고 있습니다. 연락처는 {row[3]} 입니다."
    for row in task_rows:
        db_crawls += f"'{row[0]}' 업무는 {row[1]} 상태이며 {'일정은 '+row[2] +'입니다.' if row[2] else ''} {'담당자는 ' +row[3] +'입니다.' if row[3] else ''}"
    for row in payreq_rows:
        # print(row[2] if )
        db_crawls += f"{row[1]}님에게 결재 신청한 {row[0]} 문서는 {row[2].strftime('%G-%m-%d %T')+ '에 승인되었습니다.' if row[2] else '미승인 상태 입니다.'}"
    for row in payres_rows:
        db_crawls += f"{row[1]}님이 결재 신청한 {row[0]} 문서는 {row[2].strftime('%G-%m-%d %T')+ '에 승인하셨습니다.' if row[2] else '미승인 상태 입니다.'}"

    
    db_split = text_splitter.split_text(db_crawls)
    
    return db_split

def process_file(file_path, user_input, user_id):
    text_sum = ''
    db_split = db_crawler(user_id)

    for file in file_path:
        try:
            reader = PdfReader(file)
            for page in reader.pages:  # 페이지 별로 텍스트 추출
                text = page.extract_text()
                corrected_text = text.encode('utf-8', errors='ignore').decode('utf-8')  # 인코딩 오류 무시 및 텍스트 누적
                text_sum += corrected_text + '\n'
        except Exception as e:
            print(e)
            continue
    splits = text_splitter.split_text(text_sum)
    splits += db_split
    # if len(splits) == 0: 
    #         return False
        
    # Create FAISS index
    vectorstore = FAISS.from_texts(splits, embedding_model)
    
    question = f'사용자의 입력: {user_input}'+'''
    사용자의 입력을 보고 주제를 정해서 문서 요약 관련 보고서를 만들어줘
    {
        'title': '보고서의 제목',
        'content':list['보고서의 목차별 제목'] 20글자 이내,
        'summary': '보고서의 개요' 1000글자 이내
    }
    key is title, content, summary.
    Respond using JSON only.'''
    memory = ConversationBufferMemory(
        memory_key='chat_history',
        return_messages=True,
    )

    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm_model_json,
        retriever=vectorstore.as_retriever(),
        condense_question_prompt=rag_prompt,
        memory=memory,
    )
    res = conversation_chain({'question': question})
    response = res['chat_history'][1].content.replace('\n', '').lstrip().rstrip()
    response = json.loads(response)

    return create_docx(response, vectorstore)

# Function to create the DOCX file
def create_docx(response, vectorstore):
    doc = Document()
    title = response.get('title', '제목 없음')
    doc.add_heading(title, level=0)

    doc.add_heading('목차', level=1)
    doc.add_paragraph('1. 개요')
    doc.add_paragraph('2. 본문')
    for n, cont in enumerate(response.get('content', []), 1):
        doc.add_paragraph(f'\t2-{n}. {cont}')

    doc.add_heading('1. 개요', level=1)
    doc.add_paragraph(response.get('summary', ''))

    doc.add_heading('2. 본문', level=1)
    for n, cont in enumerate(response.get('content', []), 1):
        question = f'제목이 "{title}"인 보고서의 "{cont}" 부분 상세 내용을 글로 풀어서 적어줘'    # {title}라는 보고서의 {cont} 부분 상세 내용 markdown 형식으로
        memory = ConversationBufferMemory(
            memory_key='chat_history',
            return_messages=True
        )
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=llm_model,
            retriever=vectorstore.as_retriever(),
            condense_question_prompt=rag_prompt,
            memory=memory
        )
        res = conversation_chain({'question': question})
        doc.add_heading(f'\t2-{n}. {cont}', level=2)
        doc.add_paragraph(res['chat_history'][1].content)
    
    # Format date to avoid issues in filenames
    doc.save(f'./processed/{datetime.datetime.today().strftime("%g%m%d%H%M%S")}.docx')
    return f'/processed/{datetime.datetime.today().strftime("%g%m%d%H%M%S")}.docx'

@app.route('/summary', methods=['POST'])
def summary():
    user_input = request.form.get('inputText', '').strip()
    user_id = request.form.get('user_id')
    files = [request.files[i] for i in request.files]
    processedFilePath = process_file(files, user_input, user_id)
    
    if processedFilePath:
        return jsonify({
            'result': 'ok',
            'processedFilePath': processedFilePath
        })
    else:
        return jsonify({
            'result': 'error',
            'error': '텍스트를 인식할 수 없음'
        })

@app.route('/chat', methods=['POST'])
def chat():
    params = request.get_json()
    user_input = params.get('user_input')
    user_id = params.get('user_id')
    splits = db_crawler(user_id)
    vectorstore = FAISS.from_texts(splits, embedding_model)

    memory = ConversationBufferMemory(
        memory_key='chat_history',
        return_messages=True
    )
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm_model,
        retriever=vectorstore.as_retriever(),
        condense_question_prompt=chat_prompt,
        memory=memory
    )
    
    conn = dbcp.getconn()
    cur = conn.cursor()
    
    cur.execute(
        "select ac_question, ac_answer from ai_chat_logs where user_id=%s order by chat_at asc",
        (user_id,)
    )
    rows = cur.fetchall()
    
    cur.close()
    dbcp.putconn(conn)
    
    for row in rows:
        memory.save_context(
            inputs={'human': row[0]},
            outputs={'ai': row[1]}
        )
        
    res = conversation_chain({'question': user_input})
    
    conn = dbcp.getconn()
    cur = conn.cursor()
    
    cur.execute(
        "insert into ai_chat_logs (user_id, ac_question, ac_answer) values (%s, %s, %s)",
        (user_id, user_input, res['answer'])
    )
    conn.commit()
    
    cur.close()
    dbcp.putconn(conn)
    return jsonify({'answer': res['answer']})
       
if __name__ == '__main__':
    app.run(host= host, port=port)